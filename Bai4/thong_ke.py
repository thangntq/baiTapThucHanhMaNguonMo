import pandas as pd
from numpy import array
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from tkinter import messagebox
import os

#muc 1: tong so sinh vien
def sum_sv(in_data):
    tong_sv = np.sum(in_data[:, 1])
    return tong_sv

#muc 2: tong so sinh vien dat tung loai diem
def tong_sv_type(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem = Text.split(",")
    tong_Ldiem = []
    for i in range(3, 11):
        tong_Ldiem.append(np.sum(in_data[:, i]))
    result = dict(zip(loai_diem, tong_Ldiem))
    data = result
    categories = list(data.keys())
    values = list(data.values())
    plt.figure(1, figsize=(8, 8))
    plt.bar(categories, values)
    plt.xlabel('Loại điểm')
    plt.ylabel('Số lượng')
    plt.title('Biểu đồ thể hiện số lượng từng loại điểm')
    plt.yticks(np.arange(0, 100, 5))
    plt.savefig('muc_2.png')
    #plt.show()
    plt.close()
    return ''

#muc 3.1: Phần trăm số SV đạt của môn học
def phan_tram_dat(in_data):
    part = ['sv đạt', 'sv trượt']
    rate = []
    sv_dat = []
    for i in range(1, 9):
        n = np.array(in_data[i, :])
        sv_dat.append(np.sum(n[3:10]))
    phan_tram_sv_dat = np.sum(sv_dat) / sum_sv(in_data)
    rate.append(phan_tram_sv_dat)
    rate.append(1 - phan_tram_sv_dat)
    colors = ['lightgreen', 'lightcoral']
    explode = (0, 0.1)  # Để phân cách một phần nhỏ (loại Đủ) ra xa     
    plt.figure(2, figsize=(8, 8))
    plt.pie(rate, explode=explode, labels=part, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('Số lượng sinh viên đạt và không đạt\n')
    plt.axis('equal')  # Đảm bảo biểu đồ hình tròn hoàn chỉnh
    plt.savefig('muc_3.png')
    #plt.show()
    plt.close()
    return ''

#muc 3:Tìm lớp có số SV đạt >=D nhiều nhất/ ít nhất
def Max_Min_pass_class(in_data):
    sv_dat = []
    option = ['Max', 'Min']
    lop2 = []
    for i in range(1, 9):
        n = np.array(in_data[i, :])
        sv_dat.append(np.sum(n[3:10]))
    lop1 = in_data[:, 0]
    a1 = sv_dat.index(max(sv_dat))
    lop2.append(lop1[a1])
    a2 = sv_dat.index(min(sv_dat))
    lop2.append(lop1[a2])
    Dict1 = dict(zip(option, lop2))
    return f'-- Lớp có số sinh viên đạt nhiều nhất là: {Dict1["Max"]}\n-- Lớp có số sinh viên đạt nhiều nhất là: {Dict1["Min"]}'

#muc 4: lop co nhieu/it diem A, B, C nhat
def ABC_Max_Min(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    option1= Text.split(",")
    option2 = ['Max', 'Min']
    lop1 = in_data[:, 0]
    lop3 = []
    for i in range(3, 11):
        my_dict = {}
        lop2 = []
        A = list(in_data[:, i])
        ind = A.index(max(A))
        lop2.append(lop1[ind])
        ind = A.index(min(A))
        lop2.append(lop1[ind])
        my_dict = dict(zip(option2, lop2))
        lop3.append(my_dict)
    my_dict1 = dict(zip(option1, lop3))
    df = pd.DataFrame.from_dict(my_dict1, orient='index')
    return df
    #return my_dict1

#muc 5: tim diem nao co sinh vien dat nhieu nhat
def Type_max_min(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    loai_diem2 = []
    option = ['Max', 'Min']
    tong_Ldiem = []
    for i in range(3, 11):
        tong_Ldiem.append(np.sum(in_data[:, i]))
    ind1 = tong_Ldiem.index(max(tong_Ldiem))
    loai_diem2.append(loai_diem1[ind1])
    ind2 = tong_Ldiem.index(min(tong_Ldiem))
    loai_diem2.append(loai_diem1[ind2])
    result = dict(zip(option, loai_diem2))
    return f'-- Loại điểm sinh viên đạt được nhiều nhất là: {result["Max"]}\n-- Loại điểm sinh viên đạt được ít nhất là: {result["Min"]}'

#muc 6: Tìm TBC số sv đạt điểm A,B.. của cả 9 lớp
def tbc_loai_diem(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    tbc = []
    for i in range(3, 11):
        tbc.append(round(np.mean((in_data[:, i])), 0))
    mydict = dict(zip(loai_diem1, tbc))
    data = mydict
    categories = list(data.keys())
    values = list(data.values())
    plt.figure(3, figsize=(8, 8))
    plt.bar(categories, values)
    plt.xlabel('Loại điểm')
    plt.ylabel('Số lượng sinh viên')
    plt.title('Biểu đồ thể hiện \ntrung bình sinh viên đạt từng loại điểm của 9 lớp')
    #plt.show()
    plt.savefig('muc_7.png')
    plt.close()
    return mydict

#mục 7: tổng số sinh viên đạt bài cuối kỳ
def sv_ck(in_data):
    sv_cuoi_ky = in_data[:, 15]
    return np.sum(sv_cuoi_ky)

#muc 8: so sanh so sinh vien dat chuan L1, L2
def ipL1_l2(in_data):
    sv_L1 = np.sum(in_data[:, 12])
    sv_L2 = np.sum(in_data[:, 13])
    if (sv_L1 > sv_L2) :
        return f'=> Số sinh viên đạt L1 lớn hơn L2 với L1: {sv_L1} và L2: {sv_L2}'
    elif (sv_L1 < sv_L2):
        return f'=> Số sinh viên đạt L2 lớn hơn L1 với L1: {sv_L1} và L2: {sv_L2}'
    else:
        return f'=> Số sinh viên đạt L2 bằng L1 với L1: {sv_L1} và L2: {sv_L2}'

#muc 9: do thi pho diem tung lop
def draw_graph_mark(in_data):
    Ma_lop = list(in_data[:, 0])
    Pho_diem = []
    Pd_lop = {}
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    for i in range(0, 9):
        info = list(in_data[i, :])
        type_diem = []
        for j in range(3, 11):
            type_diem.append(info[j])
        my_dict1 = dict(zip(loai_diem1, type_diem))
        Pho_diem.append(my_dict1)
    Pd_lop = dict(zip(Ma_lop, Pho_diem))
    
    data = Pd_lop
    tenlop = list(data.keys())
    Pho_diem = list(data[tenlop[0]].keys())
    x = np.arange(len(tenlop))
    bottom_values = [0] * len(tenlop)
    for loai_diem in Pho_diem:
        values = [data[lop][loai_diem] for lop in tenlop]
        plt.bar(x, values, label=loai_diem, bottom=bottom_values)
        bottom_values = [bottom + value for bottom, value in zip(bottom_values, values)]
    plt.xlabel('Lớp')
    plt.ylabel('số lượng điểm')
    plt.title('Phổ điểm của từng lớp')
    plt.xticks(x, tenlop)
    plt.yticks(np.arange(0, 65, 2))
    plt.legend()
    fig = plt.gcf()
    fig.set_size_inches(19, 8)
    plt.savefig('muc_10.png', dpi = 300) 
    #plt.show()
    plt.close()
    return ''

# muc 10: số phần trăm từng loại điểm và hiển thị bằng biểu đồ hình tròn:
def phantram_loai_diem(in_data):
    df=pd.read_csv('diemPython.csv',index_col=0,header = 0)
    diem_loai = df.loc[:, 'Loại A':'Loại F']
    tong_so_sv = diem_loai.sum() # tính tổng số sinh viên
    labels = tong_so_sv.index  # Tên loại điểm (A, B+, B, C+, C, D+, D, F)
    sizes = tong_so_sv.values  # Số lượng sinh viên đạt từng loại điểm
    colors = ['red', 'yellowgreen', 'lightblue', 'darkgreen', 'blue', 'pink', 'lightsalmon', 'brown']
    explode = (0.1, 0, 0, 0, 0, 0, 0, 0)  # Để phân cách một phần nhỏ (loại A) ra xa

    plt.figure(4, figsize=(8, 8))
    plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('Phân phối điểm theo loại')
    plt.axis('equal')  # Đảm bảo biểu đồ hình tròn hoàn chỉnh
    plt.savefig('muc_11.png')
    #plt.show()
    plt.close()
    return ''
# muc 11: Tính số phần trăm sinh viên đi thi đầy đủ
def phantram_sv_tham_gia_day_du(in_data):
    so_sv_du = np.sum(in_data[:, 15])
    so_sv_khong_du = np.sum(in_data[:, 1]) - so_sv_du

    # Tạo biểu đồ hình tròn
    labels = ['Tham gia đầy đủ', 'Tham gia không đầy đủ']
    sizes = [so_sv_du, so_sv_khong_du]
    colors = ['lightgreen', 'lightcoral']
    explode = (0.1, 0)  # Để phân cách một phần nhỏ (loại Đủ) ra xa     
    plt.figure(5, figsize=(8, 8))
    plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('Phân phối số lượng sinh viên tham gia thi KTHP\n')
    plt.axis('equal')  # Đảm bảo biểu đồ hình tròn hoàn chỉnh
    plt.savefig('muc_12.png')
    #plt.show()
    plt.close()
    return ''

#muc 13: xuất báo cáo theo từng mục
def xuat_report(in_data, teacher_name, subject, faculty, filesave, font, size, type):
    try:
        doc = Document()

        #tiêu đề
        Title = 'BÁO CÁO KẾT QUẢ KẾT THÚC HỌC PHẦN'
        title = doc.add_paragraph(Title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa tiêu đề
        
        doc.add_paragraph('\n')

        # Thông tin hoc phan

        Giang_vien = f'Giảng viên: {teacher_name}'
        Hoc_phan = f'Học phần: {subject}'
        Khoa = f'Khoa: {faculty}'

        Text1 = [Giang_vien, Hoc_phan, Khoa]
        for txt in Text1:
            title = doc.add_paragraph(txt)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề

        doc.add_paragraph('\n')

        # Thêm kết quả học tập
        title = doc.add_paragraph('A. Thống kê kết quả kết thúc học phần')
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề
        Text3 = []
        muc_1 = f'=>Mục 1: Tổng số SV tham gia môn học là: {sum_sv(in_data)} sinh viên'
        Text3.append(muc_1)
        muc_2 = f'=>Mục 2: Tổng số lượng điểm từng loại : \n{tong_sv_type(in_data)}'
        Text3.append(muc_2)
        muc_3 = f'=>Mục 3: Phần trăm số sinh viên đạt là: \n{phan_tram_dat(in_data)}'
        Text3.append(muc_3)
        muc_4 = f'=>Mục 4: Lớp có số sinh viên đạt nhiều/ít nhất: \n{Max_Min_pass_class(in_data)}'
        Text3.append(muc_4)
        muc_5 = f'=>Mục 5: Lớp có sinh viên nhiều/ít nhất ở từng loại điểm: \n{ABC_Max_Min(in_data)}'
        Text3.append(muc_5)
        muc_6 = f'=>Mục 6: Loại điểm có sinh viên đạt nhiều/ ít nhất: \n{Type_max_min(in_data)}'
        Text3.append(muc_6)
        muc_7 = f'=>Mục 7: Trung bình sinh viên đạt từng loại điểm của từng lớp: \n{tbc_loai_diem(in_data)}'
        Text3.append(muc_7)
        muc_8 = f'=>Mục 8: Tổng số sinh viên đạt bài cuối kỳ: \n--Số sinh viên đạt bài cuối kỳ là: {sv_ck(in_data)}'
        Text3.append(muc_8)
        muc_9 = f'=>Mục 9: So sánh tổng số sinh viên đạt chuẩn đầu ra L1 so với L2: \n{ipL1_l2(in_data)}'
        Text3.append(muc_9)
        muc_10 = f'=>Mục 10: Đồ thị phổ điểm của từng lớp: \n{draw_graph_mark(in_data)}'
        Text3.append(muc_10)
        muc_11 = f'=>Mục 11: Tỉ lệ phần trăm từng loại điểm: \n{phantram_loai_diem(in_data)}'
        Text3.append(muc_11)
        muc_12 = f'=>Mục 12: Tỉ lệ phần trăm số sinh viên đi thi đầy đủ: \n{phantram_sv_tham_gia_day_du(in_data)}'
        Text3.append(muc_12)
        
        for txt, i in zip(Text3, range(1, len(Text3)+1)):
            if (txt == muc_2 or txt == muc_7):
                title = doc.add_paragraph(txt)
                image_path = f'muc_{i}.png'  # Thay thế bằng đường dẫn đến tệp ảnh thực tế
                run = title.add_run()
                run.add_picture(image_path, width=Inches(6), height=Inches(5))
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề
                os.remove(f'muc_{i}.png')
            elif (txt == muc_10):
                title = doc.add_paragraph(txt)
                image_path = f'muc_{i}.png'  # Thay thế bằng đường dẫn đến tệp ảnh thực tế
                run = title.add_run()
                picture = run.add_picture(image_path, width=Inches(8), height=Inches(5))
                picture.rotation = 90
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề
                os.remove(f'muc_{i}.png')
            elif (txt == muc_3 or txt == muc_11 or txt == muc_12):
                title = doc.add_paragraph(txt)
                image_path = f'muc_{i}.png'  # Thay thế bằng đường dẫn đến tệp ảnh thực tế
                run = title.add_run()
                picture = run.add_picture(image_path, width=Inches(6), height=Inches(6))
                picture.rotation = 90
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề
                os.remove(f'muc_{i}.png')
            else:
                title = doc.add_paragraph(txt)
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Căn giữa tiêu đề
        title_style = title.style
        title_font = title_style.font
        title_font.name = font  # Tên font chữ
        title_font.size = Pt(size)  # Kích thước chữ (Pt)
        title_font.bold = True  # In đậm
        title_font.italic = False  # Nghiêng chữ
        # Lưu tài liệu vào tệp Word
        if type == 'Word':
            doc.save(f'{filesave}.docx')
            messagebox.showinfo('Thông báo','Báo cáo đã được tạo và lưu thành công.')
        else:
            messagebox.showinfo('Thông báo','Hiện tại phần mềm chưa xử lý được ở định dạng này!.')

    except PermissionError:
        messagebox.showwarning('Cảnh báo','File được lưu đang được mở.')